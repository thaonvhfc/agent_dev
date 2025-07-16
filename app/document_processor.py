import os
from typing import List
from pypdf import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from app.config import settings
import re
import unicodedata
import docx
#from doc2docx import convert

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", "; ", " ", ""]
        )
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Trích xuất văn bản từ file PDF"""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Lỗi khi đọc file PDF: {str(e)}")
    

    def starts_with_uppercase(self, line: str) -> bool:
        if not line:
            return False
        first_char = line.lstrip()[0]  # Bỏ khoảng trắng đầu dòng, lấy ký tự đầu tiên
        return unicodedata.category(first_char) == 'Lu'  # 'Lu' = Letter, uppercase

    def clean_pdf_text_lines(self, text:str):
        lines = text.split("\n")
        line_counts = {}
        for line in lines:
            line_counts[line] = line_counts.get(line, 0) + 1

        cleaned_text = [line for line in lines if line_counts[line] <= 3]  # xuất hiện quá 3 lần là nghi header/footer
        cleaned_lines = []
        for i, line in enumerate(cleaned_text):
            line = line.strip()
            # Bỏ dòng trống hoặc dòng rất ngắn nghi ngờ là header/footer
            if not line or len(line) < 5:
                continue

            # Nối dòng nếu dòng hiện tại không kết thúc bằng dấu câu và dòng tiếp theo không phải bắt đầu bằng chữ hoa
            if cleaned_lines:
                prev_line = cleaned_lines[-1]
                if not re.search(r'[.!?…]$', prev_line) and not self.starts_with_uppercase(line):
                    cleaned_lines[-1] = prev_line + ' ' + line
                    continue

            cleaned_lines.append(line)

        #print(cleaned_lines[1:10])
        return "\n".join(cleaned_lines)


    def split_text_into_chunks(self, text: str, source: str) -> List[Document]:
        """Chia văn bản thành các chunk nhỏ"""
        chunks = self.text_splitter.split_text(text)
        documents = []
        
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "source": source,
                    "chunk_id": i,
                    "total_chunks": len(chunks)
                }
            )
            documents.append(doc)
        
        return documents
    
    def process_pdf(self, file_path: str) -> List[Document]:
        """Xử lý file PDF và trả về danh sách documents"""
        text = self.extract_text_from_pdf(file_path)
        text = self.clean_pdf_text_lines(text)  # làm sạch văn bản từ PDF
        #print(f"Text đã trích xuất từ file:\n\n{text}...")
        filename = os.path.basename(file_path)
        documents = self.split_text_into_chunks(text, filename)
        return documents
    
    # Xử lý các file khác nhau như DOCX, TXT, PDF
    def extract_text_from_file(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            text = self.extract_text_from_pdf(file_path)
            text = self.clean_pdf_text_lines(text)  # làm sạch văn bản từ PDF
            return text
        elif ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif ext in ['.doc', '.docx']:
            return self.extract_text_from_docx(file_path)
        else:
            raise Exception("Định dạng file không hỗ trợ")

    def extract_text_from_docx(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])

    def process_file(self, file_path: str) -> List[Document]:
        text = self.extract_text_from_file(file_path)
        filename = os.path.basename(file_path)
        documents = self.split_text_into_chunks(text, filename)
        return documents
    
    